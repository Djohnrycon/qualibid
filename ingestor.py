"""File ingestor for QualiBid. Reads contractor submissions and returns text."""
from pathlib import Path
import pdfplumber
from openpyxl import load_workbook
from docx import Document


def read_pdf(path):
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
    return "\n\n".join(parts)


def read_excel(path):
    wb = load_workbook(path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
        parts.append("")
    return "\n".join(parts)


def read_docx(path):
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def read_image(path):
    return f"[Image file: {Path(path).name} - OCR not implemented in demo]"


def read_file(path):
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        return read_pdf(p)
    if ext == ".xlsx":
        return read_excel(p)
    if ext == ".docx":
        return read_docx(p)
    if ext in (".png", ".jpg", ".jpeg"):
        return read_image(p)
    return f"[Unsupported file type: {ext}]"
